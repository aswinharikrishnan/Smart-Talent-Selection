from docx import Document

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    combined_text = []

    # 1. Grab every paragraph in the main document body
    for para in doc.paragraphs:
        if para.text.strip():
            combined_text.append(para.text)

    # 2. Grab every table (Most resumes put 'Experience' in tables)
    for table in doc.tables:
        for row in table.rows:
            # Join all cells in a row to maintain context
            row_content = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_content:
                combined_text.append(" | ".join(row_content))

    return "\n".join(combined_text)